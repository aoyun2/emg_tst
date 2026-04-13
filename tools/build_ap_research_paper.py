from __future__ import annotations

import re
from pathlib import Path

import bibtexparser
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
PAPER_MD = ROOT / 'docs' / 'ap_research_paper_revised.md'
RAW_BIB = ROOT / 'artifacts' / 'current_paper_sources.bib'
OUT_BIB = ROOT / 'docs' / 'ap_research_current_project_sources.bib'
OUT_RUBRIC = ROOT / 'docs' / 'ap_research_rubric_alignment.md'
OUT_DOCX = Path(r'C:\Users\aaron\Downloads\Research Paper — Aaron (Current Project).docx')
AP_RUBRIC_PDF = ROOT / 'artifacts' / 'ap25_research_scoring_guidelines.pdf'

KEY_REMAP = {
    'DILLINGHAM_2002': 'Dillingham2002',
    'Farina_2014': 'Farina2014',
    'Frisch_1933': 'FrischWaugh1933',
    'Hargrove_2007': 'Hargrove2007',
    'Huang_2008': 'Huang2008',
    'Hudgins_1993': 'Hudgins1993',
    'Huynh_2009': 'Huynh2009',
    'Kele__2023': 'Keles2023',
    'Krasoulis_2019': 'Krasoulis2019',
    'Hof_2005': 'Hof2005',
    'Li_2023': 'Li2023RVR',
    'Liang_2023': 'Liang2023',
    'Lin_2024': 'LinHe2024',
    'Lin_2025': 'LinZhangZhao2025',
    'Lovell_1963': 'Lovell1963',
    'Phinyomark_2018': 'Phinyomark2018',
    'Pinhey_2022': 'Pinhey2022',
    'Pran_2021': 'Pran2021',
    'Roberts_2017': 'Roberts2017',
    'Spearman_1904': 'Spearman1904',
    'Todorov_2012': 'Todorov2012',
    'Tougui_2021': 'Tougui2021',
    'Zerveas_2021': 'Zerveas2021',
    'Zhang_2022': 'Zhang2022',
    'Zhu_2022': 'Zhu2022',
    'Ahkami_2023': 'Ahkami2023',
    'Cimolato_2022': 'Cimolato2022',
    'Camargo_2021': 'Camargo2021',
    'Sun_2022': 'Sun2022',
    'Moghadam_2023': 'Moghadam2023',
    'Hur_2025': 'Hur2025',
    'Hof_2008': 'Hof2008',
    'Kim_2015': 'Kim2015',
    'Huber_1964': 'Huber1964',
    'Hochreiter_1997': 'HochreiterSchmidhuber1997',
    'Schuster_1997': 'SchusterPaliwal1997',
    'De_Luca_1997': 'DeLuca1997',
}

MANUAL_ENTRIES = r'''
@article{Pickle2018,
  title = {The Inverted Pendulum Model Is Insufficient to Explain the Reactive Stepping Strategies Used in Bipedal Walking on a Known Slippery Surface},
  author = {Pickle, Nathan T. and Wilken, Jason M. and Aldridge Whitehead, Joanna M. and Silverman, Anne K.},
  journal = {Journal of Biomechanics},
  year = {2018},
  volume = {77},
  pages = {176--183},
  doi = {10.1016/j.jbiomech.2018.06.024}
}

@article{Gill2019,
  title = {The Utility of the Extrapolated Center of Mass in the Assessment of Human Balance after Foot Placement},
  author = {Gill, Stephanie V. and Narain, Alyssa and Arora, Harsh and Smith, Courtney and Bhamidipati, Pranav},
  journal = {Journal of Biomechanics},
  year = {2019},
  volume = {97},
  pages = {109356},
  doi = {10.1016/j.jbiomech.2019.109356}
}
'''


def clean_bib_text(raw: str) -> str:
    text = raw
    li_tag = '@article{Li_2023,'
    if text.count(li_tag) >= 2:
        first = text.find(li_tag)
        second = text.find(li_tag, first + len(li_tag))
        text = text[:first] + '@article{Li2023RVR,' + text[first + len(li_tag):second] + '@article{Li2023MultiTaskTransformer,' + text[second + len(li_tag):]
    for old, new in KEY_REMAP.items():
        text = text.replace(f'{{{old},', f'{{{new},')
    if '@article{Pickle2018,' not in text:
        text = text.strip() + '\n\n' + MANUAL_ENTRIES.strip() + '\n'
    return text


def _strip_braces(s: str) -> str:
    return re.sub(r'[{}]', '', (s or '')).strip()


def _authors_list(author_field: str) -> list[str]:
    return [_strip_braces(a).strip() for a in (author_field or '').replace('\n', ' ').split(' and ') if a.strip()]


def _format_author_name(name: str) -> str:
    name = ' '.join(name.split())
    if ',' in name:
        last, first = [p.strip() for p in name.split(',', 1)]
    else:
        parts = name.split()
        last = parts[-1]
        first = ' '.join(parts[:-1])
    initials = ' '.join(f'{p[0]}.' for p in re.split(r'[\s-]+', first) if p)
    return f'{last}, {initials}'.strip().replace(' ,', ',')


def format_authors(author_field: str) -> str:
    authors = [_format_author_name(a) for a in _authors_list(author_field)]
    if not authors:
        return ''
    if len(authors) == 1:
        return authors[0]
    if len(authors) == 2:
        return f'{authors[0]}, & {authors[1]}'
    return ', '.join(authors[:-1]) + f', & {authors[-1]}'


def format_reference(entry: dict) -> str:
    typ = entry.get('ENTRYTYPE', 'article')
    authors = format_authors(entry.get('author', ''))
    year = _strip_braces(entry.get('year', 'n.d.'))
    title = _strip_braces(entry.get('title', ''))
    doi = _strip_braces(entry.get('doi', ''))
    url = _strip_braces(entry.get('url', ''))
    pages = _strip_braces(entry.get('pages', '')).replace('--', '–')
    volume = _strip_braces(entry.get('volume', ''))
    number = _strip_braces(entry.get('number', ''))
    journal = _strip_braces(entry.get('journal', ''))
    booktitle = _strip_braces(entry.get('booktitle', ''))
    publisher = _strip_braces(entry.get('publisher', ''))
    howpublished = _strip_braces(entry.get('howpublished', ''))
    eprint = _strip_braces(entry.get('eprint', ''))
    archive = _strip_braces(entry.get('archivePrefix', ''))

    if typ == 'article':
        middle = journal
        if volume and number:
            middle += f', {volume}({number})'
        elif volume:
            middle += f', {volume}'
        ref = f'{authors} ({year}). {title}. {middle}'
        if pages:
            ref += f', {pages}'
        ref += '.'
        if doi:
            ref += f' https://doi.org/{doi}'
        elif url:
            ref += f' {url}'
        return ref
    if typ == 'inproceedings':
        ref = f'{authors} ({year}). {title}. In {booktitle}'
        if pages:
            ref += f' (pp. {pages})'
        ref += '.'
        if publisher:
            ref += f' {publisher}.'
        if doi:
            ref += f' https://doi.org/{doi}'
        elif url:
            ref += f' {url}'
        return ref
    ref = f'{authors} ({year}). {title}.'
    if archive and eprint:
        ref += f' {archive}:{eprint}.'
    if howpublished:
        ref += f' {howpublished}.'
    if url:
        ref += f' {url}'
    return ref


def generate_bib() -> list[dict]:
    raw = RAW_BIB.read_text(encoding='utf-8')
    cleaned = clean_bib_text(raw)
    OUT_BIB.write_text(cleaned, encoding='utf-8')
    db = bibtexparser.loads(cleaned)
    entries = list(db.entries)
    entries.sort(key=lambda e: (_authors_list(e.get('author', ''))[0].split(',')[0].strip().lower() if _authors_list(e.get('author', '')) else 'zzz', _strip_braces(e.get('year', ''))))
    return entries


def strip_markup(text: str) -> str:
    return text.replace('**', '').replace('`', '')


def count_words(text: str) -> int:
    main = text.split('## References')[0]
    main = re.sub(r'\[\[FIGURE:.*?\]\]', ' ', main)
    main = main.replace('**', '').replace('`', '')
    return len(re.findall(r"[A-Za-z0-9°\-]+", main))


def write_rubric_alignment(word_count: int) -> None:
    text = f'''# AP Research Rubric Alignment

Official rubric source used for this rewrite: [ap25_research_scoring_guidelines.pdf]({AP_RUBRIC_PDF.as_posix()})

## Goal

This revision was written to align with the official AP Research academic paper criteria for a top-score submission. The paper was rebuilt around the current Georgia Tech + CNN-BiLSTM + MoCapAct/XCoM project rather than the older custom-data transformer version.

## Alignment to Score-5 Expectations

### 1. Focused and justified inquiry
- The paper now asks one narrow question: whether lower wearable knee-angle RMSE corresponds to lower excess instability in paired simulation after motion-match quality is controlled.
- The gap is explicit: knee-angle models are usually evaluated with offline accuracy metrics, but those metrics do not necessarily reveal whole-body physical usefulness.

### 2. Detailed and replicable method
- The rewritten methods section defines the dataset, signal channels, preprocessing, windowing, forecast horizon, architecture, optimizer, loss, validation strategy, simulation setup, instability metric, and statistics in enough detail to reproduce the benchmark.
- The paper explicitly states that the benchmark is leave-one-file-out, not subject-holdout.
- The paper defines excess instability AUC, motion-match controls, and the rank-based Frisch-Waugh-Lovell partial Spearman procedure step by step.

### 3. Rich analysis leading to a new understanding
- The results section reports both the strong prediction benchmark and the null partial-correlation result.
- The discussion does not overclaim. Instead, it explains the specific new understanding supported by the evidence: sub-10° predictive accuracy did not independently correspond to lower excess instability in this benchmark.
- Limitations and future work are stated explicitly rather than hidden.

### 4. Clear communication
- The writing style follows the stronger early-introduction tone of the original paper rather than the later AI-generated sections.
- Definitions are embedded where the terms first appear.
- The manuscript is organized around the AP expectation that claims, method choices, and interpretations must be justified rather than merely asserted.

### 5. Citation discipline
- In-text citations are used frequently after claims, methodological decisions, and interpretive statements that require support.
- The revised package includes both a textual reference list in the paper and a machine-verifiable BibTeX file with more than 40 entries.

## Current Draft Word Count

- Main body word count estimate: **{word_count}**
- This count excludes the references list and figure captions.
'''
    OUT_RUBRIC.write_text(text, encoding='utf-8')


def render_docx(paper_text: str, references: list[str]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal.font.size = Pt(12)

    for style_name in ['Title', 'Heading 1', 'Heading 2', 'Heading 3']:
        st = doc.styles[style_name]
        st.font.name = 'Times New Roman'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    for raw in paper_text.splitlines():
        line = raw.rstrip()
        if not line:
            doc.add_paragraph('')
            continue
        mfig = re.match(r'\[\[FIGURE:(.+?)\|(.*?)\]\]', line)
        if mfig:
            rel_path, caption = mfig.groups()
            p = (ROOT / rel_path).resolve()
            if p.exists():
                pic = doc.add_paragraph()
                pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic.add_run().add_picture(str(p), width=Inches(6.1))
                cap = doc.add_paragraph(strip_markup(caption))
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.line_spacing = 1.0
            continue
        if line.startswith('# '):
            p = doc.add_paragraph(style='Title')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(strip_markup(line[2:]))
            continue
        if line.startswith('## '):
            doc.add_paragraph(strip_markup(line[3:]), style='Heading 1')
            continue
        if line.startswith('### '):
            doc.add_paragraph(strip_markup(line[4:]), style='Heading 2')
            continue
        if re.match(r'^\d+\.\s', line):
            doc.add_paragraph(strip_markup(line), style='List Number')
            continue
        if line.startswith('- '):
            doc.add_paragraph(strip_markup(line[2:]), style='List Bullet')
            continue
        p = doc.add_paragraph(strip_markup(line))
        p.paragraph_format.line_spacing = 2.0

    doc.add_page_break()
    doc.add_paragraph('References', style='Heading 1')
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.line_spacing = 2.0

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))


def main() -> None:
    refs = generate_bib()
    ref_lines = [format_reference(e) for e in refs]
    paper = PAPER_MD.read_text(encoding='utf-8')
    if '## References' in paper:
        paper = paper.split('## References')[0].rstrip() + '\n\n## References\n\n[[REFERENCES]]\n'
    wc = count_words(paper)
    paper = re.sub(r'\*\*Word Count:\*\*\s*`?\[\[WORD_COUNT\]\]`?|\*\*Word Count:\*\*\s*\d+', f'**Word Count:** {wc}', paper)
    paper = paper.replace('[[REFERENCES]]', '\n\n'.join(ref_lines))
    PAPER_MD.write_text(paper, encoding='utf-8')
    write_rubric_alignment(wc)
    render_docx(paper, ref_lines)
    print(f'word_count={wc}')
    print(f'docx={OUT_DOCX}')
    print(f'bib={OUT_BIB}')
    print(f'rubric={OUT_RUBRIC}')


if __name__ == '__main__':
    main()
