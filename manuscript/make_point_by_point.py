"""Generate the point-by-point response to the technical check as a Word document."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

MANUSCRIPT_ID = "57f1a801-3acf-47ad-b878-69cde9d135f8"

OPENING = (
    "I thank the editorial office for the technical check. Both points have been "
    "addressed, and the responses are set out below. Each editorial comment is "
    "reproduced in italics and is followed by my response."
)

POINTS = [
    ("Please provide City in the affiliation for author(s) in the manuscript.",
     "The city has been added to the author affiliation. The affiliation now reads "
     "\"Spring Branch Academic Institute, Houston, Texas, United States\" and appears "
     "in this form on the title page of the manuscript and in the affiliation "
     "metadata of the submission."),

    ("Please provide the cover letter in word .doc format.",
     "A cover letter has been prepared in Microsoft Word format and is uploaded with "
     "this submission as cover_letter.docx."),
]

CLOSING = (
    "No other changes were made to the manuscript in response to the technical "
    "check. I would be glad to supply any further information the editorial office "
    "requires."
)

SIGNATURE = ["Yours sincerely,", "", "Aaron Xiong",
             "Spring Branch Academic Institute", "Houston, Texas, United States",
             "aaxiong2008@gmail.com"]


def build(out_path: Path) -> Path:
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(10)
    style.paragraph_format.line_spacing = 1.15

    for section in document.sections:
        section.left_margin = section.right_margin = Inches(1.0)
        section.top_margin = section.bottom_margin = Inches(1.0)

    heading = document.add_paragraph()
    run = heading.add_run("Response to Technical Check")
    run.bold = True
    run.font.size = Pt(13)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    reference = document.add_paragraph()
    reference.add_run(f"Manuscript ID: {MANUSCRIPT_ID} v1.0")
    reference.paragraph_format.space_after = Pt(14)

    document.add_paragraph(OPENING)

    for index, (comment, response) in enumerate(POINTS, start=1):
        block = document.add_paragraph()
        label = block.add_run(f"Point {index}. ")
        label.bold = True
        remark = block.add_run(comment)
        remark.italic = True
        block.paragraph_format.space_after = Pt(4)

        answer = document.add_paragraph()
        prefix = answer.add_run("Response: ")
        prefix.bold = True
        answer.add_run(response)
        answer.paragraph_format.left_indent = Inches(0.25)

    document.add_paragraph(CLOSING)

    document.add_paragraph()
    for line in SIGNATURE:
        paragraph = document.add_paragraph(line)
        paragraph.paragraph_format.space_after = Pt(0)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(out_path)
    return out_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--out", type=Path,
                        default=Path("manuscript/point_by_point_response.docx"))
    args = parser.parse_args()
    path = build(args.out)
    print(f"wrote {path} ({path.stat().st_size / 1024:.0f} KB)")
