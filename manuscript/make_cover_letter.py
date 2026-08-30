"""Generate the submission cover letter as a Word document.

Kept as a script so the letter can be regenerated if the reported figures change,
in the same way the manuscript itself is generated from the run artifacts.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pathlib
from docx.shared import Inches, Pt

import json
import os

_RUNS = pathlib.Path(os.environ.get("EMG_TST_RUNS", r"C:\Users\aaron\emg_data\runs"))
_ACC = json.loads(
    (_RUNS / "analysis" / "checkpoint_correlation.json").read_text(encoding="utf-8")
)["accuracy_level"]
_SPLIT = f"{_ACC['breakpoint_rmse_deg']:.2f}"
_LO, _HI = (f"{v:.2f}" for v in _ACC["breakpoint_95pct_ci"])
_ABOVE = f"{_ACC['above_breakpoint']['slope_per_degree']:+.5f}"
_BELOW = f"{_ACC['below_breakpoint']['slope_per_degree']:+.5f}"

TITLE = ("Toward the use of simulated environments to evaluate sEMG-inf"
         "ormed knee-angle prediction: RMSE predicts simulated "
         "instability only above a threshold accuracy")

BODY = [
    "Dear Editor,",

    "I am submitting the enclosed manuscript, titled \"" + TITLE + ",\" for "
    "consideration as an original research article.",

    "Root-mean-square error is the standard reported metric for continuous "
    "knee-angle prediction, and it is generally treated as though a lower value "
    "implies a more useful controller. That assumption has not been tested "
    "directly, because a study that evaluates a single converged model can only "
    "compare prediction windows against one another, and differences between "
    "windows are dominated by factors unrelated to the predictor.",

    "The present study addresses this by holding the evaluation panel fixed and "
    "varying the model instead. Eighty one-second walking windows were matched to "
    "MoCapAct reference motions and replayed in paired MuJoCo simulations at "
    "fourteen accuracy levels sampled along the model's own gradient-descent "
    "training path, giving 1,120 rollouts over an accuracy range from 22 to 4.9 "
    "degrees. In each rollout the predicted knee angle was the tracking target of "
    "a proportional-derivative override on the right knee, and the paired "
    "reference condition ran the unmodified expert policy on the same matched "
    "motion, so that the two conditions differed only in the target given to the "
    "knee.",

    "The relationship between prediction error and simulated instability was found "
    "to be non-monotonic. An exploratory two-segment split falls at " + _SPLIT + " degrees "
    "root-mean-square error (95% confidence interval " + _LO + " to " + _HI + "). Above that "
    "accuracy, worse prediction produced more simulated instability, with a slope "
    "of " + _ABOVE + " seconds per degree. Below it the relationship inverted, with "
    "a slope of " + _BELOW + " seconds per degree, which is consistent with a partially "
    "fitted model commanding a flatter trajectory than the recorded one. "
    "Both slopes were estimated by resampling participants and carrying each across "
    "every accuracy level, and both intervals exclude zero.",

    "This result has a practical consequence for how prosthetic regression models "
    "are evaluated. A converged predictor operates in the range where further "
    "reduction in root-mean-square error no longer predicts better simulated "
    "behavior, and a null result obtained there should be expected rather than "
    "interpreted as evidence that accuracy is unimportant. The surface "
    "electromyography ablation reported in the manuscript serves as a quality gate "
    "confirming that the predictor was worth evaluating physically, rather than as "
    "a contribution in itself; mean participant error was 4.748 degrees with the "
    "correction and 5.093 degrees without it, and 70 of 90 participants improved.",

    "The manuscript has not been published elsewhere and is not under "
    "consideration by any other journal. The study used only publicly available "
    "datasets and did not involve new human or animal participants, so ethical "
    "approval was not required. I am the sole author, I declare no competing "
    "interests, and I received no funding for this work. The analysis code, the "
    "unattended run pipeline, and the script that generates every reported figure "
    "and numerical value are available in the project repository.",

    "Thank you for your consideration.",
]

CLOSING = ["Yours sincerely,", "", "Aaron Xiong",
           "Spring Branch Academic Institute", "Houston, Texas, United States",
           "aaxiong2008@gmail.com"]


def build(out_path: Path) -> Path:
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(10)
    style.paragraph_format.line_spacing = 1.15

    for section in document.sections:
        section.left_margin = section.right_margin = Inches(1.0)
        section.top_margin = section.bottom_margin = Inches(1.0)

    for text in BODY:
        paragraph = document.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_paragraph()
    for line in CLOSING:
        paragraph = document.add_paragraph(line)
        paragraph.paragraph_format.space_after = Pt(0)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(out_path)
    return out_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--out", type=Path,
                        default=Path("manuscript/cover_letter.docx"))
    args = parser.parse_args()
    path = build(args.out)
    print(f"wrote {path} ({path.stat().st_size / 1024:.0f} KB)")
